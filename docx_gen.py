import shutil
from docx import Document
from config import TEMPLATE_PATH


def make_initial(fish: str) -> str:
    parts = fish.strip().split()
    if len(parts) >= 3:
        return f"{parts[1][0]}.{parts[2][0]}.{parts[0]}"
    elif len(parts) == 2:
        return f"{parts[1][0]}.{parts[0]}"
    return fish


def _replace_para(para, reps: dict):
    full = "".join(r.text for r in para.runs)
    if not any(k in full for k in reps):
        return
    new = full
    for old, val in reps.items():
        new = new.replace(old, val)
    if para.runs:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""


def generate(data: dict) -> str:
    fish    = data["full_name"]
    guruh   = data["guruh"]
    initial = make_initial(fish)

    out = f"ariza_{fish.replace(' ', '_')}.docx"
    shutil.copy(TEMPLATE_PATH, out)
    doc = Document(out)

    reps = {
        "Iqtisodiyot va axborot texnologiyalari fakulteti": data["fakultet"],
        "Iqtisodiyot yo'nalishi": data["yonalish"],
        "Toshtemirov Dilmurod Xasanovich": fish,
        "IQT-4-23-guruh": f"{guruh}-guruh",
        "IQT-4-23": guruh,
        "4-23-guruh": f"{guruh}-guruh",
        "4-23": guruh,
        "D.X.Toshtemirov": initial,
    }

    for para in doc.paragraphs:
        _replace_para(para, reps)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para, reps)

    doc.save(out)
    return out
